"""
Document text and metadata extraction services.
Handles OCR, text extraction, and NLP-based metadata parsing.
"""
import io
import re
import json
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime, date
from dataclasses import dataclass, field, asdict
from enum import Enum
import logging

import pdfplumber
from PIL import Image
import pytesseract
import docx
import openpyxl
from dateutil import parser as date_parser
import spacy
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.storage import get_storage_client
from app.models.document import Document

logger = logging.getLogger(__name__)


class ExtractionMethod(str, Enum):
    """Extraction methods used."""
    PDFPLUMBER = "pdfplumber"
    TESSERACT_OCR = "tesseract_ocr"
    PYTHON_DOCX = "python-docx"
    OPENPYXL = "openpyxl"
    DIRECT = "direct"


@dataclass
class ExtractedEntity:
    """Represents an extracted entity from text."""
    type: str
    text: str
    value: Any = None
    confidence: float = 1.0
    position: Optional[int] = None


@dataclass
class ContractParty:
    """Represents a party in a contract."""
    name: str
    role: Optional[str] = None
    entity_type: Optional[str] = None  # corporation, individual, etc.


@dataclass
class MonetaryValue:
    """Represents a monetary amount in contract."""
    amount: float
    currency: str
    context: Optional[str] = None
    position: Optional[int] = None


@dataclass
class ContractClause:
    """Represents a contract clause."""
    title: str
    text: str
    section_number: Optional[str] = None
    position: Optional[int] = None


@dataclass
class ImportantDate:
    """Represents an important date in contract."""
    text: str
    date_value: Optional[date] = None
    context: Optional[str] = None
    position: Optional[int] = None


@dataclass
class DocumentMetadata:
    """Extracted document metadata."""
    title: Optional[str] = None
    document_type: Optional[str] = None
    contract_number: Optional[str] = None
    parties: List[ContractParty] = field(default_factory=list)
    effective_date: Optional[date] = None
    expiration_date: Optional[date] = None
    important_dates: List[ImportantDate] = field(default_factory=list)
    monetary_values: List[MonetaryValue] = field(default_factory=list)
    clauses: List[ContractClause] = field(default_factory=list)
    governing_law: Optional[str] = None
    jurisdiction: Optional[str] = None
    reference_numbers: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage."""
        data = {}
        for key, value in asdict(self).items():
            if value is not None:
                if isinstance(value, date):
                    data[key] = value.isoformat()
                elif isinstance(value, list):
                    data[key] = [
                        item.isoformat() if isinstance(item, date) else 
                        asdict(item) if hasattr(item, '__dataclass_fields__') else 
                        item 
                        for item in value
                    ]
                else:
                    data[key] = value
        return data


@dataclass
class TextExtractionResult:
    """Result of text extraction operation."""
    success: bool
    text: str = ""
    page_count: int = 0
    extraction_method: Optional[str] = None
    ocr_confidence: Optional[float] = None
    error_message: Optional[str] = None


@dataclass
class ExtractionResult:
    """Complete extraction result including text and metadata."""
    success: bool
    extracted_text: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    page_count: int = 0
    extraction_method: Optional[str] = None
    processing_time: float = 0.0
    error_message: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return asdict(self)


class TextExtractor:
    """Extracts text from various document formats."""
    
    def __init__(self):
        """Initialize text extractor."""
        self.supported_formats = {
            'application/pdf': self.extract_from_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.extract_from_docx,
            'application/msword': self.extract_from_doc,
            'text/plain': self.extract_from_txt,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self.extract_from_xlsx,
            'application/vnd.ms-excel': self.extract_from_xls,
        }
    
    async def extract(self, content: bytes, mime_type: str, **kwargs) -> TextExtractionResult:
        """Extract text from document based on mime type."""
        extractor = self.supported_formats.get(mime_type)
        
        if not extractor:
            return TextExtractionResult(
                success=False,
                error_message=f"Unsupported file format: {mime_type}"
            )
        
        try:
            return await extractor(content, **kwargs)
        except Exception as e:
            logger.error(f"Extraction failed: {str(e)}")
            return TextExtractionResult(
                success=False,
                error_message=str(e)
            )
    
    async def extract_from_pdf(self, content: bytes, use_ocr: bool = False) -> TextExtractionResult:
        """Extract text from PDF, with OCR support for scanned documents."""
        try:
            text_parts = []
            page_count = 0
            
            # Try regular text extraction first
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            # If no text extracted and OCR requested, use OCR
            if not text_parts and use_ocr:
                return await self._extract_with_ocr(content)
            
            # Check if we got meaningful text
            full_text = "\n".join(text_parts)
            if len(full_text.strip()) < 50 and use_ocr:
                # Very little text, might be scanned
                return await self._extract_with_ocr(content)
            
            return TextExtractionResult(
                success=True,
                text=full_text,
                page_count=page_count,
                extraction_method=ExtractionMethod.PDFPLUMBER
            )
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return TextExtractionResult(
                success=False,
                error_message=f"PDF extraction failed: {str(e)}"
            )
    
    async def _extract_with_ocr(self, pdf_content: bytes) -> TextExtractionResult:
        """Extract text using OCR from scanned PDF."""
        try:
            import pdf2image
            
            # Convert PDF to images
            images = pdf2image.convert_from_bytes(pdf_content)
            
            text_parts = []
            confidence_scores = []
            
            for img in images:
                # Perform OCR
                ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                
                # Extract text with confidence
                page_text = []
                page_confidence = []
                
                for i, text in enumerate(ocr_data['text']):
                    if text.strip():
                        page_text.append(text)
                        page_confidence.append(ocr_data['conf'][i])
                
                text_parts.append(" ".join(page_text))
                if page_confidence:
                    confidence_scores.extend(page_confidence)
            
            # Calculate average confidence
            avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
            
            return TextExtractionResult(
                success=True,
                text="\n".join(text_parts),
                page_count=len(images),
                extraction_method=ExtractionMethod.TESSERACT_OCR,
                ocr_confidence=avg_confidence / 100.0  # Convert to 0-1 scale
            )
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return TextExtractionResult(
                success=False,
                error_message=f"OCR extraction failed: {str(e)}"
            )
    
    async def extract_from_docx(self, content: bytes) -> TextExtractionResult:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(io.BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return TextExtractionResult(
                success=True,
                text="\n".join(text_parts),
                page_count=1,  # DOCX doesn't have pages in the same way
                extraction_method=ExtractionMethod.PYTHON_DOCX
            )
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return TextExtractionResult(
                success=False,
                error_message=f"DOCX extraction failed: {str(e)}"
            )
    
    async def extract_from_doc(self, content: bytes) -> TextExtractionResult:
        """Extract text from DOC file (legacy format)."""
        # For legacy DOC files, we might need to use python-docx2txt or similar
        # For now, return unsupported
        return TextExtractionResult(
            success=False,
            error_message="Legacy DOC format not yet supported"
        )
    
    async def extract_from_txt(self, content: bytes) -> TextExtractionResult:
        """Extract text from plain text file."""
        try:
            text = content.decode('utf-8', errors='ignore')
            return TextExtractionResult(
                success=True,
                text=text,
                page_count=1,
                extraction_method=ExtractionMethod.DIRECT
            )
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            return TextExtractionResult(
                success=False,
                error_message=f"TXT extraction failed: {str(e)}"
            )
    
    async def extract_from_xlsx(self, content: bytes) -> TextExtractionResult:
        """Extract text from Excel file."""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            
            for sheet in workbook.worksheets:
                sheet_text = [f"Sheet: {sheet.title}"]
                
                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    
                    if row_values:
                        sheet_text.append("\t".join(row_values))
                
                text_parts.append("\n".join(sheet_text))
            
            return TextExtractionResult(
                success=True,
                text="\n\n".join(text_parts),
                page_count=len(workbook.worksheets),
                extraction_method=ExtractionMethod.OPENPYXL
            )
            
        except Exception as e:
            logger.error(f"XLSX extraction failed: {str(e)}")
            return TextExtractionResult(
                success=False,
                error_message=f"XLSX extraction failed: {str(e)}"
            )
    
    async def extract_from_xls(self, content: bytes) -> TextExtractionResult:
        """Extract text from legacy Excel file."""
        # Would need xlrd or similar for legacy XLS
        return TextExtractionResult(
            success=False,
            error_message="Legacy XLS format not yet supported"
        )


class MetadataExtractor:
    """Extracts structured metadata from contract text using NLP."""
    
    def __init__(self):
        """Initialize metadata extractor with NLP model."""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except:
            # If model not installed, use basic extraction
            self.nlp = None
            logger.warning("Spacy model not loaded, using basic extraction")
    
    async def extract_metadata(self, text: str) -> DocumentMetadata:
        """Extract structured metadata from contract text."""
        metadata = DocumentMetadata()
        
        # Extract title
        metadata.title = self._extract_title(text)
        metadata.document_type = self._extract_document_type(text)
        
        # Extract contract number
        metadata.contract_number = self._extract_contract_number(text)
        metadata.reference_numbers = self._extract_reference_numbers(text)
        
        # Extract parties
        metadata.parties = self._extract_parties(text)
        
        # Extract dates
        dates = self._extract_dates(text)
        metadata.effective_date = dates.get('effective_date')
        metadata.expiration_date = dates.get('expiration_date')
        metadata.important_dates = dates.get('other_dates', [])
        
        # Extract monetary values
        metadata.monetary_values = self._extract_monetary_values(text)
        
        # Extract clauses
        metadata.clauses = self._extract_clauses(text)
        
        # Extract governing law and jurisdiction
        legal_info = self._extract_legal_jurisdiction(text)
        metadata.governing_law = legal_info.get('governing_law')
        metadata.jurisdiction = legal_info.get('jurisdiction')
        
        return metadata
    
    def _extract_title(self, text: str) -> Optional[str]:
        """Extract document title."""
        lines = text.split('\n')
        
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and len(line) < 100:
                # Check if line is in caps or looks like a title
                if line.isupper() or re.match(r'^[A-Z][A-Z\s]+$', line):
                    return line
                # Check for common title patterns
                if re.match(r'^(CONTRACT|AGREEMENT|MEMORANDUM|TERMS)', line, re.I):
                    return line
        
        return None
    
    def _extract_document_type(self, text: str) -> Optional[str]:
        """Extract document type."""
        doc_types = [
            'Service Agreement', 'Purchase Agreement', 'License Agreement',
            'Employment Contract', 'Non-Disclosure Agreement', 'NDA',
            'Master Service Agreement', 'MSA', 'Statement of Work', 'SOW',
            'Terms and Conditions', 'Terms of Service', 'Privacy Policy'
        ]
        
        text_lower = text.lower()
        for doc_type in doc_types:
            if doc_type.lower() in text_lower:
                return doc_type
        
        return None
    
    def _extract_contract_number(self, text: str) -> Optional[str]:
        """Extract contract number."""
        patterns = [
            r'Contract\s+(?:Number|No\.?|#)\s*:?\s*([A-Z0-9\-]+)',
            r'Agreement\s+(?:Number|No\.?|#)\s*:?\s*([A-Z0-9\-]+)',
            r'Reference\s*:?\s*([A-Z0-9\-]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.I)
            if match:
                return match.group(1)
        
        return None
    
    def _extract_reference_numbers(self, text: str) -> List[str]:
        """Extract reference numbers."""
        references = []
        
        # Look for PO numbers, reference numbers, etc.
        patterns = [
            r'PO\s*(?:Number|No\.?|#)?\s*:?\s*([A-Z0-9\-]+)',
            r'Purchase\s+Order\s*:?\s*([A-Z0-9\-]+)',
            r'Reference\s*:?\s*([A-Z0-9\-]+)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.I)
            references.extend(matches)
        
        return list(set(references))  # Remove duplicates
    
    def _extract_parties(self, text: str) -> List[ContractParty]:
        """Extract contract parties."""
        parties = []
        
        # Pattern for party identification
        patterns = [
            r'between\s+([^,\n]+?)\s*\(["\']?(\w+)["\']?\)',
            r'between\s+([^,\n]+?)\s+and\s+([^,\n]+)',
            r'Party\s+[AB]\s*:\s*([^\n]+)',
            r'\(["\']?(\w+)["\']?\)\s*:\s*([^\n]+)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.I | re.M)
            for match in matches:
                if isinstance(match, tuple):
                    if len(match) == 2:
                        name = match[0].strip()
                        role = match[1].strip() if len(match[1]) < 20 else None
                        if name:
                            parties.append(ContractParty(name=name, role=role))
                else:
                    parties.append(ContractParty(name=match.strip()))
        
        # Use NER if available
        if self.nlp and len(parties) < 2:
            doc = self.nlp(text[:5000])  # Analyze first 5000 chars
            for ent in doc.ents:
                if ent.label_ == "ORG":
                    parties.append(ContractParty(name=ent.text, entity_type="organization"))
        
        return parties[:5]  # Return max 5 parties
    
    def _extract_dates(self, text: str) -> Dict[str, Any]:
        """Extract important dates."""
        dates = {
            'effective_date': None,
            'expiration_date': None,
            'other_dates': []
        }
        
        # Patterns for specific date types
        effective_patterns = [
            r'Effective\s+Date\s*:?\s*([^\n]+)',
            r'dated\s+as\s+of\s+([^\n]+)',
            r'commenc\w+\s+on\s+([^\n]+)',
        ]
        
        expiration_patterns = [
            r'Expir\w+\s+Date\s*:?\s*([^\n]+)',
            r'terminat\w+\s+on\s+([^\n]+)',
            r'continue\s+until\s+([^\n]+)',
            r'end\w+\s+on\s+([^\n]+)',
        ]
        
        # Extract effective date
        for pattern in effective_patterns:
            match = re.search(pattern, text, re.I)
            if match:
                try:
                    parsed_date = date_parser.parse(match.group(1), fuzzy=True)
                    dates['effective_date'] = parsed_date.date()
                    break
                except:
                    pass
        
        # Extract expiration date
        for pattern in expiration_patterns:
            match = re.search(pattern, text, re.I)
            if match:
                try:
                    parsed_date = date_parser.parse(match.group(1), fuzzy=True)
                    dates['expiration_date'] = parsed_date.date()
                    break
                except:
                    pass
        
        # Extract other dates
        date_pattern = r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        date_matches = re.findall(date_pattern, text)
        
        for date_str in date_matches[:10]:  # Limit to 10 dates
            try:
                parsed_date = date_parser.parse(date_str)
                dates['other_dates'].append(ImportantDate(
                    text=date_str,
                    date_value=parsed_date.date()
                ))
            except:
                pass
        
        return dates
    
    def _extract_monetary_values(self, text: str) -> List[MonetaryValue]:
        """Extract monetary amounts."""
        values = []
        
        # Patterns for monetary values
        patterns = [
            r'\$\s*([\d,]+\.?\d*)',
            r'USD\s*([\d,]+\.?\d*)',
            r'([\d,]+\.?\d*)\s*(?:dollars|USD)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.I)
            for match in matches:
                try:
                    amount = float(match.replace(',', ''))
                    if amount > 0:
                        values.append(MonetaryValue(
                            amount=amount,
                            currency="USD"
                        ))
                except:
                    pass
        
        return values[:10]  # Return max 10 values
    
    def _extract_clauses(self, text: str) -> List[ContractClause]:
        """Extract contract clauses."""
        clauses = []
        
        # Pattern for numbered sections
        section_pattern = r'(\d+\.?\s*)([A-Z][A-Z\s]+)\n([^\n]+(?:\n[^\d][^\n]*)*)'
        
        matches = re.findall(section_pattern, text)
        for match in matches[:20]:  # Limit to 20 clauses
            section_num = match[0].strip()
            title = match[1].strip()
            content = match[2].strip()
            
            if len(title) < 50:  # Reasonable title length
                clauses.append(ContractClause(
                    title=title,
                    text=content[:500],  # First 500 chars
                    section_number=section_num
                ))
        
        return clauses
    
    def _extract_legal_jurisdiction(self, text: str) -> Dict[str, Optional[str]]:
        """Extract governing law and jurisdiction."""
        result = {
            'governing_law': None,
            'jurisdiction': None
        }
        
        # Governing law patterns
        law_patterns = [
            r'governed\s+by\s+the\s+laws?\s+of\s+(?:the\s+)?([^,\n]+)',
            r'under\s+the\s+laws?\s+of\s+(?:the\s+)?([^,\n]+)',
        ]
        
        for pattern in law_patterns:
            match = re.search(pattern, text, re.I)
            if match:
                result['governing_law'] = match.group(1).strip()
                break
        
        # Jurisdiction patterns
        jurisdiction_patterns = [
            r'courts?\s+of\s+([^,\n]+)',
            r'jurisdiction\s+of\s+([^,\n]+)',
        ]
        
        for pattern in jurisdiction_patterns:
            match = re.search(pattern, text, re.I)
            if match:
                result['jurisdiction'] = match.group(1).strip()
                break
        
        return result


class ExtractionPipeline:
    """Complete extraction pipeline for documents."""
    
    def __init__(self, db_session: Optional[AsyncSession] = None, min_ocr_confidence: float = 0.0):
        """Initialize extraction pipeline."""
        self.db_session = db_session
        self.text_extractor = TextExtractor()
        self.metadata_extractor = MetadataExtractor()
        self.min_ocr_confidence = min_ocr_confidence
    
    async def download_document(self, document: Document) -> bytes:
        """Download document from storage."""
        storage = await get_storage_client()
        
        # Download from MinIO
        import io
        from app.core.storage import download_file
        
        file_stream = await download_file(
            storage,
            "legal-documents",
            document.file_path
        )
        
        if file_stream:
            return file_stream.read()
        else:
            raise Exception(f"Failed to download document: {document.file_path}")
    
    async def process_document(self, document: Document) -> ExtractionResult:
        """Process a single document through the extraction pipeline."""
        import time
        start_time = time.time()
        
        try:
            # Download document
            content = await self.download_document(document)
            
            # Extract text
            text_result = await self.text_extractor.extract(
                content,
                document.mime_type,
                use_ocr=True
            )
            
            if not text_result.success:
                # Update document status
                if self.db_session:
                    document.extraction_status = "failed"
                    document.extraction_error = text_result.error_message
                    await self.db_session.commit()
                
                return ExtractionResult(
                    success=False,
                    error_message=text_result.error_message,
                    processing_time=time.time() - start_time
                )
            
            # Check OCR confidence if applicable
            if text_result.extraction_method == ExtractionMethod.TESSERACT_OCR:
                if text_result.ocr_confidence and text_result.ocr_confidence < self.min_ocr_confidence:
                    error_msg = f"OCR confidence too low: {text_result.ocr_confidence:.2f}"
                    
                    if self.db_session:
                        document.extraction_status = "failed"
                        document.extraction_error = error_msg
                        await self.db_session.commit()
                    
                    return ExtractionResult(
                        success=False,
                        error_message=error_msg,
                        processing_time=time.time() - start_time
                    )
            
            # Extract metadata
            metadata = await self.metadata_extractor.extract_metadata(text_result.text)
            
            # Update document in database
            if self.db_session:
                document.extracted_text = text_result.text
                document.extraction_status = "completed"
                document.extraction_error = None
                document.metadata = {
                    **(document.metadata or {}),
                    **metadata.to_dict(),
                    'page_count': text_result.page_count,
                    'extraction_method': text_result.extraction_method,
                    'ocr_confidence': text_result.ocr_confidence
                }
                await self.db_session.commit()
            
            return ExtractionResult(
                success=True,
                extracted_text=text_result.text,
                metadata=metadata.to_dict(),
                page_count=text_result.page_count,
                extraction_method=text_result.extraction_method,
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            logger.error(f"Pipeline processing failed: {str(e)}")
            
            if self.db_session:
                document.extraction_status = "failed"
                document.extraction_error = str(e)
                await self.db_session.commit()
            
            return ExtractionResult(
                success=False,
                error_message=str(e),
                processing_time=time.time() - start_time
            )
    
    async def process_batch(self, documents: List[Document]) -> List[ExtractionResult]:
        """Process multiple documents in batch."""
        results = []
        
        for document in documents:
            result = await self.process_document(document)
            results.append(result)
        
        return results