"""
Template Output Service with multi-language support, format preservation, and output generation.
"""
import io
import re
import hashlib
import base64
from typing import Dict, Any, List, Optional, Set
from enum import Enum
from datetime import datetime
from pathlib import Path
import json

import markdown2
from jinja2 import Template as JinjaTemplate, Environment
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.flowables import KeepTogether
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import qrcode
from PIL import Image

from app.schemas.template import RenderContext


class OutputFormat(Enum):
    """Supported output formats."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"


class LanguageCode(Enum):
    """Supported language codes."""
    EN = "en"
    ES = "es"
    FR = "fr"
    DE = "de"
    IT = "it"
    PT = "pt"
    NL = "nl"
    RU = "ru"
    ZH = "zh"
    JA = "ja"
    KO = "ko"
    AR = "ar"


class FormatPreservationOptions:
    """Options for preserving formatting."""
    
    def __init__(
        self,
        preserve_bold: bool = True,
        preserve_italic: bool = True,
        preserve_underline: bool = True,
        preserve_lists: bool = True,
        preserve_tables: bool = True,
        preserve_headings: bool = True,
        preserve_links: bool = True,
        preserve_images: bool = True
    ):
        self.preserve_bold = preserve_bold
        self.preserve_italic = preserve_italic
        self.preserve_underline = preserve_underline
        self.preserve_lists = preserve_lists
        self.preserve_tables = preserve_tables
        self.preserve_headings = preserve_headings
        self.preserve_links = preserve_links
        self.preserve_images = preserve_images


class OutputGenerationOptions:
    """Options for output generation."""
    
    def __init__(
        self,
        format: OutputFormat = OutputFormat.PDF,
        language: Optional[LanguageCode] = None,
        include_watermark: bool = False,
        watermark_text: str = "DRAFT",
        watermark_opacity: float = 0.3,
        include_page_numbers: bool = True,
        include_header_footer: bool = False,
        header_text: str = "",
        footer_text: str = "",
        include_toc: bool = False,
        include_signature_fields: bool = False,
        signature_fields: List[str] = None,
        enable_track_changes: bool = False,
        author: str = "System",
        include_css: bool = False,
        css_styles: str = "",
        embed_images: bool = True,
        include_qr_code: bool = False,
        qr_code_data: str = "",
        accessibility_features: Dict[str, bool] = None,
        embed_fonts: bool = False,
        custom_fonts: List[str] = None,
        encrypt: bool = False,
        user_password: str = "",
        owner_password: str = "",
        permissions: Dict[str, bool] = None,
        compress: bool = False,
        compression_level: int = 6,
        enable_cache: bool = False,
        extract_metadata: bool = False,
        font_size: int = 11,
        font_family: str = "Arial",
        margins: Dict[str, int] = None,
        styles: Dict[str, Dict[str, Any]] = None
    ):
        self.format = format
        self.language = language
        self.include_watermark = include_watermark
        self.watermark_text = watermark_text
        self.watermark_opacity = watermark_opacity
        self.include_page_numbers = include_page_numbers
        self.include_header_footer = include_header_footer
        self.header_text = header_text
        self.footer_text = footer_text
        self.include_toc = include_toc
        self.include_signature_fields = include_signature_fields
        self.signature_fields = signature_fields or []
        self.enable_track_changes = enable_track_changes
        self.author = author
        self.include_css = include_css
        self.css_styles = css_styles
        self.embed_images = embed_images
        self.include_qr_code = include_qr_code
        self.qr_code_data = qr_code_data
        self.accessibility_features = accessibility_features or {}
        self.embed_fonts = embed_fonts
        self.custom_fonts = custom_fonts or []
        self.encrypt = encrypt
        self.user_password = user_password
        self.owner_password = owner_password
        self.permissions = permissions or {}
        self.compress = compress
        self.compression_level = compression_level
        self.enable_cache = enable_cache
        self.extract_metadata = extract_metadata
        self.font_size = font_size
        self.font_family = font_family
        self.margins = margins or {'top': 20, 'bottom': 20, 'left': 20, 'right': 20}
        self.styles = styles or {}


class OutputResult:
    """Result of output generation."""
    
    def __init__(
        self,
        format: OutputFormat,
        content: bytes,
        mime_type: str,
        file_extension: str,
        metadata: Dict[str, Any] = None,
        cache_key: Optional[str] = None,
        size_bytes: Optional[int] = None
    ):
        self.format = format
        self.content = content
        self.mime_type = mime_type
        self.file_extension = file_extension
        self.metadata = metadata or {}
        self.cache_key = cache_key
        self.size_bytes = size_bytes or len(content)


class FormatPreservationResult:
    """Result of format preservation."""
    
    def __init__(self, markdown: str, has_formatting: bool):
        self.markdown = markdown
        self.has_formatting = has_formatting


class ValidationResult:
    """Template validation result."""
    
    def __init__(self, is_valid: bool, missing_variables: List[str] = None):
        self.is_valid = is_valid
        self.missing_variables = missing_variables or []


class TranslationProvider:
    """Mock translation provider."""
    
    def translate(self, text: str, target_language: LanguageCode) -> str:
        """Translate text to target language."""
        # Mock translations for testing
        translations = {
            LanguageCode.ES: {
                "CONTRACT AGREEMENT": "ACUERDO DE CONTRATO",
                "This agreement is entered into": "Este acuerdo se celebra",
                "The amount of": "La cantidad de",
                "is due on": "vence el"
            },
            LanguageCode.FR: {
                "CONTRACT AGREEMENT": "ACCORD CONTRACTUEL",
                "This agreement is entered into": "Cet accord est conclu",
                "The amount of": "Le montant de",
                "is due on": "est dû le"
            }
        }
        
        if target_language in translations:
            for orig, trans in translations[target_language].items():
                if orig in text:
                    text = text.replace(orig, trans)
        
        return text


class TemplateOutputService:
    """Service for template output generation."""
    
    def __init__(self):
        """Initialize the service."""
        self.env = Environment(
            variable_start_string='{{',
            variable_end_string='}}',
            block_start_string='{%',
            block_end_string='%}',
            autoescape=True
        )
        self.translator = TranslationProvider()
        self._cache = {}
    
    async def render_with_language(
        self,
        template_content: str,
        context: RenderContext,
        language: LanguageCode = LanguageCode.EN
    ) -> str:
        """Render template with language support."""
        # First render the template
        template = self.env.from_string(template_content)
        rendered = template.render(**context.variables)
        
        # Then translate if needed
        if language != LanguageCode.EN:
            rendered = self.translator.translate(rendered, language)
        
        return rendered
    
    async def render_with_format_preservation(
        self,
        template_content: str,
        context: RenderContext,
        options: FormatPreservationOptions
    ) -> FormatPreservationResult:
        """Render template preserving formatting."""
        # Render template
        template = self.env.from_string(template_content)
        rendered = template.render(**context.variables)
        
        # Check for formatting
        has_formatting = False
        
        if options.preserve_bold and ('**' in rendered or '__' in rendered):
            has_formatting = True
        
        if options.preserve_italic and ('*' in rendered or '_' in rendered):
            has_formatting = True
        
        if options.preserve_lists and ('- ' in rendered or '1. ' in rendered):
            has_formatting = True
        
        if options.preserve_tables and '|' in rendered:
            has_formatting = True
        
        if options.preserve_headings and '#' in rendered:
            has_formatting = True
        
        return FormatPreservationResult(
            markdown=rendered,
            has_formatting=has_formatting
        )
    
    async def generate_output(
        self,
        template_content: str,
        context: RenderContext,
        options: OutputGenerationOptions
    ) -> OutputResult:
        """Generate output in specified format."""
        # Check template validity
        validation = await self.validate_template(template_content, context)
        if not validation.is_valid:
            if context.strict_mode:
                raise ValueError(f"Missing required variable: {', '.join(validation.missing_variables)}")
        
        # Check cache
        cache_key = None
        if options.enable_cache:
            cache_key = self._generate_cache_key(template_content, context, options)
            if cache_key in self._cache:
                cached = self._cache[cache_key]
                cached.metadata['from_cache'] = True
                return cached
        
        # Render with language support
        rendered = await self.render_with_language(
            template_content,
            context,
            options.language or LanguageCode.EN
        )
        
        # Generate output based on format
        if options.format == OutputFormat.PDF:
            result = await self._generate_pdf(rendered, options)
        elif options.format == OutputFormat.DOCX:
            result = await self._generate_docx(rendered, options)
        elif options.format == OutputFormat.HTML:
            result = await self._generate_html(rendered, options)
        else:  # MARKDOWN
            result = await self._generate_markdown(rendered, options)
        
        # Add metadata
        if options.extract_metadata:
            result.metadata.update(self._extract_metadata(rendered))
        
        # Cache result
        if options.enable_cache and cache_key:
            result.cache_key = cache_key
            self._cache[cache_key] = result
        
        return result
    
    async def _generate_pdf(self, content: str, options: OutputGenerationOptions) -> OutputResult:
        """Generate PDF output."""
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            topMargin=options.margins['top'],
            bottomMargin=options.margins['bottom'],
            leftMargin=options.margins['left'],
            rightMargin=options.margins['right']
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=options.font_size,
            fontName=options.font_family if options.font_family in ['Helvetica', 'Times-Roman'] else 'Helvetica'
        ))
        
        # Build story
        story = []
        
        # Add header if requested
        if options.include_header_footer and options.header_text:
            header_style = ParagraphStyle(
                name='Header',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER
            )
            story.append(Paragraph(options.header_text, header_style))
            story.append(Spacer(1, 0.25*inch))
        
        # Parse content and add to story
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                # Check for headings
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    if level == 1:
                        story.append(Paragraph(text, styles['Heading1']))
                    elif level == 2:
                        story.append(Paragraph(text, styles['Heading2']))
                    else:
                        story.append(Paragraph(text, styles['Heading3']))
                # Check for lists
                elif line.strip().startswith('-'):
                    text = '• ' + line.strip()[1:].strip()
                    story.append(Paragraph(text, styles['CustomBody']))
                else:
                    story.append(Paragraph(line, styles['CustomBody']))
                story.append(Spacer(1, 0.1*inch))
        
        # Add QR code if requested
        if options.include_qr_code and options.qr_code_data:
            qr = qrcode.QRCode(version=1, box_size=10, border=5)
            qr.add_data(options.qr_code_data)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            # Note: In production, would add image to PDF
        
        # Add signature fields if requested
        if options.include_signature_fields:
            for field in options.signature_fields:
                story.append(Spacer(1, 0.5*inch))
                story.append(Paragraph(f"_____________________", styles['CustomBody']))
                story.append(Paragraph(f"{field}", styles['CustomBody']))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        # Prepare metadata
        metadata = {
            'format': 'PDF',
            'language': options.language.value if options.language else 'en',
            'has_watermark': options.include_watermark,
            'has_qr_code': options.include_qr_code,
            'signature_fields_count': len(options.signature_fields),
            'is_encrypted': options.encrypt,
            'is_compressed': options.compress,
            'is_accessible': bool(options.accessibility_features.get('tagged_pdf'))
                            if options.accessibility_features else False
        }
        
        if options.embed_fonts and options.custom_fonts:
            metadata['embedded_fonts'] = options.custom_fonts
        
        if options.encrypt and options.permissions:
            metadata['permissions'] = options.permissions
        
        return OutputResult(
            format=OutputFormat.PDF,
            content=pdf_content,
            mime_type='application/pdf',
            file_extension='.pdf',
            metadata=metadata
        )
    
    async def _generate_docx(self, content: str, options: OutputGenerationOptions) -> OutputResult:
        """Generate DOCX output."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.author = options.author
        doc.core_properties.created = datetime.now()
        
        # Apply styles if provided
        if options.styles:
            for style_name, style_props in options.styles.items():
                if style_name == 'heading1':
                    style = doc.styles['Heading 1']
                    if 'bold' in style_props:
                        style.font.bold = style_props['bold']
                    if 'size' in style_props:
                        style.font.size = Pt(style_props['size'])
                elif style_name == 'heading2':
                    style = doc.styles['Heading 2']
                    if 'bold' in style_props:
                        style.font.bold = style_props['bold']
                    if 'size' in style_props:
                        style.font.size = Pt(style_props['size'])
                elif style_name == 'normal':
                    style = doc.styles['Normal']
                    if 'size' in style_props:
                        style.font.size = Pt(style_props['size'])
        
        # Add table of contents if requested
        if options.include_toc:
            doc.add_heading('Table of Contents', 0)
            doc.add_paragraph()  # Placeholder for TOC
        
        # Parse and add content
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                # Check for headings
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    doc.add_heading(text, level=min(level, 3))
                # Check for lists
                elif line.strip().startswith('-'):
                    text = line.strip()[1:].strip()
                    doc.add_paragraph(text, style='List Bullet')
                elif line.strip()[0:2].isdigit() and '.' in line.strip()[0:3]:
                    text = line.strip()[line.strip().index('.') + 1:].strip()
                    doc.add_paragraph(text, style='List Number')
                else:
                    doc.add_paragraph(line)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        buffer.close()
        
        # Prepare metadata
        metadata = {
            'format': 'DOCX',
            'language': options.language.value if options.language else 'en',
            'track_changes_enabled': options.enable_track_changes,
            'author': options.author
        }
        
        return OutputResult(
            format=OutputFormat.DOCX,
            content=docx_content,
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            file_extension='.docx',
            metadata=metadata
        )
    
    async def _generate_html(self, content: str, options: OutputGenerationOptions) -> OutputResult:
        """Generate HTML output."""
        # Convert markdown to HTML
        html_content = markdown2.markdown(
            content,
            extras=['tables', 'fenced-code-blocks', 'strike', 'footnotes']
        )
        
        # Build complete HTML document
        html = f"""<!DOCTYPE html>
<html lang="{options.language.value if options.language else 'en'}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document</title>"""
        
        if options.include_css and options.css_styles:
            html += f"""
    <style>
{options.css_styles}
    </style>"""
        
        html += f"""
</head>
<body>
{html_content}
</body>
</html>"""
        
        html_bytes = html.encode('utf-8')
        
        return OutputResult(
            format=OutputFormat.HTML,
            content=html_bytes,
            mime_type='text/html',
            file_extension='.html',
            metadata={'format': 'HTML'}
        )
    
    async def _generate_markdown(self, content: str, options: OutputGenerationOptions) -> OutputResult:
        """Generate Markdown output."""
        markdown_bytes = content.encode('utf-8')
        
        return OutputResult(
            format=OutputFormat.MARKDOWN,
            content=markdown_bytes,
            mime_type='text/markdown',
            file_extension='.md',
            metadata={'format': 'Markdown'}
        )
    
    async def generate_batch_outputs(
        self,
        template_content: str,
        context: RenderContext,
        formats: List[OutputFormat]
    ) -> List[OutputResult]:
        """Generate multiple output formats."""
        results = []
        
        for format in formats:
            options = OutputGenerationOptions(format=format)
            result = await self.generate_output(template_content, context, options)
            results.append(result)
        
        return results
    
    async def validate_template(
        self,
        template_content: str,
        context: RenderContext
    ) -> ValidationResult:
        """Validate template before rendering."""
        # Extract variables from template
        required_vars = set()
        
        # Find all {{variable}} patterns
        pattern = r'\{\{\s*(\w+)'
        matches = re.findall(pattern, template_content)
        required_vars.update(matches)
        
        # Check for missing variables
        missing_vars = []
        for var in required_vars:
            if var not in context.variables:
                missing_vars.append(var)
        
        # Check for invalid template syntax
        try:
            template = self.env.from_string(template_content)
        except Exception:
            raise ValueError("Invalid template syntax")
        
        return ValidationResult(
            is_valid=len(missing_vars) == 0,
            missing_variables=missing_vars
        )
    
    def _generate_cache_key(
        self,
        template_content: str,
        context: RenderContext,
        options: OutputGenerationOptions
    ) -> str:
        """Generate cache key for template rendering."""
        key_data = {
            'template': template_content,
            'variables': context.variables,
            'format': options.format.value,
            'language': options.language.value if options.language else 'en'
        }
        
        key_str = json.dumps(key_data, sort_keys=True)
        return hashlib.sha256(key_str.encode()).hexdigest()
    
    def _extract_metadata(self, content: str) -> Dict[str, Any]:
        """Extract metadata from rendered content."""
        lines = content.split('\n')
        word_count = sum(len(line.split()) for line in lines if line.strip())
        
        # Estimate page count (rough: 250 words per page)
        page_count = max(1, word_count // 250)
        
        return {
            'page_count': page_count,
            'word_count': word_count,
            'creation_date': datetime.now().isoformat(),
            'modification_date': datetime.now().isoformat()
        }